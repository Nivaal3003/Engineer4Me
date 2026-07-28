"""Bounded PDF and Office parsing for Engineer4Me document ingestion.

This parser extends the existing OCR-aware parser chain with native support
for PDF, DOCX, XLSX, and legacy XLS files. It applies explicit resource
limits, validates OOXML archives before opening them, rejects encrypted or
macro-enabled inputs, and uses OCR only for PDF pages without native text.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date, datetime, time, timedelta
from decimal import Decimal
from io import BytesIO
import math
from pathlib import Path, PurePosixPath
import re
import xml.etree.ElementTree as ElementTree
from typing import Any, Sequence
from zipfile import BadZipFile, ZipFile

from defusedxml import ElementTree as SafeElementTree
from defusedxml.common import DefusedXmlException
from docx import Document as load_docx_document
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
import pypdfium2 as pdfium
import xlrd
from xlrd.biffh import error_text_from_code

from app.ingestion.document_models import (
    BoundingBox,
    ContentBlockType,
    DocumentFormat,
    DocumentUpload,
    ExtractionMethod,
    ParsedContentBlock,
    ParsedDocument,
    ParsedPage,
    ParsedTable,
    SpreadsheetCellRange,
)
from app.ingestion.document_parser import (
    DocumentParser,
    DocumentParserError,
    DocumentTooLargeError,
    EmptyDocumentError,
    MalformedDocumentError,
    UnsupportedDocumentFormatError,
    _derive_title_from_blocks,
    _extract_text_blocks,
    _normalise_document_text,
    _normalise_inline_text,
    _table_to_text,
)
from app.ingestion.ocr_document_parser import (
    OcrAwareDocumentParser,
    OcrAwareDocumentParserConfig,
)


_PDF_OFFICE_FORMATS = frozenset(
    {
        DocumentFormat.PDF,
        DocumentFormat.DOCX,
        DocumentFormat.XLSX,
        DocumentFormat.XLS,
    }
)
_PDF_OFFICE_SUFFIX_FORMATS = {
    ".pdf": DocumentFormat.PDF,
    ".docx": DocumentFormat.DOCX,
    ".xlsx": DocumentFormat.XLSX,
    ".xls": DocumentFormat.XLS,
}
_OOXML_REQUIRED_MEMBERS = {
    DocumentFormat.DOCX: frozenset(
        {
            "[Content_Types].xml",
            "word/document.xml",
        }
    ),
    DocumentFormat.XLSX: frozenset(
        {
            "[Content_Types].xml",
            "xl/workbook.xml",
        }
    ),
}
_OOXML_XML_SUFFIXES = (".xml", ".rels")
_OLE_COMPOUND_FILE_SIGNATURE = bytes.fromhex("D0CF11E0A1B11AE1")
_PDF_HEADER = b"%PDF-"
_SAFETY_PREFIX_PATTERN = re.compile(
    r"^\s*(danger|warning|caution|note)\s*[:\-]\s*(.+?)\s*$",
    re.IGNORECASE,
)
_HEADING_STYLE_PATTERN = re.compile(
    r"^heading\s+([1-9]\d*)$",
    re.IGNORECASE,
)


class PasswordProtectedDocumentError(DocumentParserError):
    """Raised when encrypted input cannot be processed safely."""


class UnsafeOfficeArchiveError(MalformedDocumentError):
    """Raised when an OOXML archive violates a safety rule."""


@dataclass(frozen=True, slots=True)
class PdfOfficeDocumentParserConfig:
    """Immutable resource controls for PDF and Office parsing."""

    fallback_parser: OcrAwareDocumentParserConfig = field(
        default_factory=OcrAwareDocumentParserConfig,
    )
    maximum_document_bytes: int = 25 * 1024 * 1024
    maximum_extracted_characters: int = 10_000_000
    maximum_pdf_pages: int = 500
    minimum_pdf_native_characters: int = 1
    enable_pdf_ocr_fallback: bool = True
    pdf_ocr_scale: float = 2.0
    maximum_pdf_ocr_page_pixels: int = 25_000_000
    maximum_pdf_ocr_total_pixels: int = 100_000_000
    maximum_archive_members: int = 10_000
    maximum_archive_uncompressed_bytes: int = 250 * 1024 * 1024
    maximum_archive_compression_ratio: float = 200.0
    maximum_docx_blocks: int = 100_000
    maximum_table_rows: int = 100_000
    maximum_table_columns: int = 1_000
    maximum_workbook_sheets: int = 256
    maximum_spreadsheet_rows: int = 100_000
    maximum_spreadsheet_columns: int = 16_384
    maximum_spreadsheet_cells: int = 2_000_000
    maximum_cell_characters: int = 32_767

    def __post_init__(self) -> None:
        if not isinstance(
            self.fallback_parser,
            OcrAwareDocumentParserConfig,
        ):
            raise TypeError(
                "fallback_parser must be an "
                "OcrAwareDocumentParserConfig."
            )

        positive_integer_fields = (
            "maximum_document_bytes",
            "maximum_extracted_characters",
            "maximum_pdf_pages",
            "minimum_pdf_native_characters",
            "maximum_pdf_ocr_page_pixels",
            "maximum_pdf_ocr_total_pixels",
            "maximum_archive_members",
            "maximum_archive_uncompressed_bytes",
            "maximum_docx_blocks",
            "maximum_table_rows",
            "maximum_table_columns",
            "maximum_workbook_sheets",
            "maximum_spreadsheet_rows",
            "maximum_spreadsheet_columns",
            "maximum_spreadsheet_cells",
            "maximum_cell_characters",
        )

        for field_name in positive_integer_fields:
            value = getattr(self, field_name)
            if (
                not isinstance(value, int)
                or isinstance(value, bool)
                or value <= 0
            ):
                raise ValueError(
                    f"{field_name} must be a positive integer."
                )

        for field_name in (
            "pdf_ocr_scale",
            "maximum_archive_compression_ratio",
        ):
            value = getattr(self, field_name)
            if (
                not isinstance(value, (int, float))
                or isinstance(value, bool)
                or not math.isfinite(float(value))
                or float(value) <= 0
            ):
                raise ValueError(
                    f"{field_name} must be a positive finite number."
                )

        if not isinstance(self.enable_pdf_ocr_fallback, bool):
            raise TypeError(
                "enable_pdf_ocr_fallback must be a boolean."
            )

        if (
            self.maximum_pdf_ocr_total_pixels
            < self.maximum_pdf_ocr_page_pixels
        ):
            raise ValueError(
                "maximum_pdf_ocr_total_pixels cannot be less than "
                "maximum_pdf_ocr_page_pixels."
            )

        object.__setattr__(
            self,
            "pdf_ocr_scale",
            float(self.pdf_ocr_scale),
        )
        object.__setattr__(
            self,
            "maximum_archive_compression_ratio",
            float(self.maximum_archive_compression_ratio),
        )


class PdfOfficeDocumentParser(DocumentParser):
    """Parse PDF and Office documents and delegate existing formats."""

    PARSER_NAME = "engineer4me-pdf-office-document-parser"
    PARSER_VERSION = "1.0.0"

    def __init__(
        self,
        config: PdfOfficeDocumentParserConfig | None = None,
    ) -> None:
        self.config = config or PdfOfficeDocumentParserConfig()
        self._fallback_parser = OcrAwareDocumentParser(
            self.config.fallback_parser,
        )

    @property
    def parser_name(self) -> str:
        """Return the stable parser identifier."""

        return self.PARSER_NAME

    @property
    def parser_version(self) -> str:
        """Return the parser contract version."""

        return self.PARSER_VERSION

    @property
    def fallback_parser(self) -> OcrAwareDocumentParser:
        """Return the parser used for existing native and image formats."""

        return self._fallback_parser

    def supports(self, upload: DocumentUpload) -> bool:
        """Return whether this parser chain supports an upload."""

        effective_format = self.resolve_document_format(upload)
        return (
            effective_format in _PDF_OFFICE_FORMATS
            or self._fallback_parser.supports(upload)
        )

    def resolve_document_format(
        self,
        upload: DocumentUpload,
    ) -> DocumentFormat:
        """Resolve PDF and Office suffixes before fallback parser rules."""

        if upload.document_format != DocumentFormat.UNKNOWN:
            return upload.document_format

        suffix = Path(upload.filename).suffix.lower()
        office_format = _PDF_OFFICE_SUFFIX_FORMATS.get(suffix)

        if office_format is not None:
            return office_format

        return self._fallback_parser.resolve_document_format(upload)

    def parse(
        self,
        upload: DocumentUpload,
        content: bytes,
    ) -> ParsedDocument:
        """Parse one upload with bounded PDF, Office, or fallback logic."""

        effective_format = self.resolve_document_format(upload)

        if effective_format not in _PDF_OFFICE_FORMATS:
            return self._fallback_parser.parse(upload, content)

        self._validate_input(upload, content)

        if upload.password_protected:
            raise PasswordProtectedDocumentError(
                f"Document '{upload.filename}' is marked as "
                "password-protected and cannot be parsed."
            )

        if effective_format == DocumentFormat.PDF:
            return self._parse_pdf(upload, content)

        if effective_format == DocumentFormat.DOCX:
            return self._parse_docx(upload, content)

        if effective_format == DocumentFormat.XLSX:
            return self._parse_xlsx(upload, content)

        if effective_format == DocumentFormat.XLS:
            return self._parse_xls(upload, content)

        raise UnsupportedDocumentFormatError(
            f"Unsupported document format: {effective_format.value}."
        )

    def _validate_input(
        self,
        upload: DocumentUpload,
        content: bytes,
    ) -> None:
        if not isinstance(content, bytes):
            raise TypeError("Document content must be supplied as bytes.")

        if not content:
            raise EmptyDocumentError(
                f"Document '{upload.filename}' contains no data."
            )

        if len(content) > self.config.maximum_document_bytes:
            raise DocumentTooLargeError(
                f"Document '{upload.filename}' is {len(content)} bytes, "
                "which exceeds the configured limit of "
                f"{self.config.maximum_document_bytes} bytes."
            )

    def _parse_pdf(
        self,
        upload: DocumentUpload,
        content: bytes,
    ) -> ParsedDocument:
        if _PDF_HEADER not in content[:1024]:
            raise MalformedDocumentError(
                f"PDF '{upload.filename}' does not contain a valid "
                "PDF header."
            )

        pages: list[ParsedPage] = []
        warnings = self._size_mismatch_warnings(upload, content)
        native_page_count = 0
        ocr_page_count = 0
        blank_page_count = 0
        total_ocr_pixels = 0
        extracted_character_count = 0
        next_sequence_number = 1
        metadata: dict[str, Any] = {}
        attachment_count = 0
        pdf_document: Any = None

        try:
            pdf_document = pdfium.PdfDocument(content)
            page_count = len(pdf_document)

            if page_count == 0:
                raise EmptyDocumentError(
                    f"PDF '{upload.filename}' contains no pages."
                )

            if page_count > self.config.maximum_pdf_pages:
                raise DocumentTooLargeError(
                    f"PDF '{upload.filename}' contains {page_count} "
                    "pages, which exceeds the configured limit of "
                    f"{self.config.maximum_pdf_pages}."
                )

            try:
                metadata = {
                    str(key): _metadata_value(value)
                    for key, value in pdf_document.get_metadata_dict(
                        skip_empty=True,
                    ).items()
                }
            except Exception:
                warnings.append(
                    "PDF metadata could not be read and was skipped."
                )

            try:
                attachment_count = int(
                    pdf_document.count_attachments()
                )
            except Exception:
                warnings.append(
                    "PDF attachment metadata could not be read."
                )

            for page_index in range(page_count):
                page_number = page_index + 1
                pdf_page: Any = None

                try:
                    pdf_page = pdf_document[page_index]
                    width, height = self._validated_pdf_page_size(
                        pdf_page,
                        page_number=page_number,
                    )
                    native_text = self._extract_pdf_page_text(
                        pdf_page,
                        page_number=page_number,
                        warnings=warnings,
                    )

                    if (
                        len(native_text)
                        >= self.config.minimum_pdf_native_characters
                    ):
                        page, next_sequence_number = (
                            self._build_native_pdf_page(
                                text=native_text,
                                page_number=page_number,
                                width=width,
                                height=height,
                                sequence_number=next_sequence_number,
                            )
                        )
                        native_page_count += 1
                    elif self.config.enable_pdf_ocr_fallback:
                        (
                            page,
                            next_sequence_number,
                            rendered_pixels,
                        ) = self._build_ocr_pdf_page(
                            upload=upload,
                            pdf_page=pdf_page,
                            page_number=page_number,
                            width=width,
                            height=height,
                            sequence_number=next_sequence_number,
                            current_ocr_pixels=total_ocr_pixels,
                        )
                        total_ocr_pixels += rendered_pixels

                        if page.text:
                            ocr_page_count += 1
                        else:
                            blank_page_count += 1
                    else:
                        page = ParsedPage(
                            page_number=page_number,
                            width=width,
                            height=height,
                            text="",
                            blocks=[],
                            extraction_method=ExtractionMethod.UNKNOWN,
                            extraction_confidence=0.0,
                            warnings=[
                                "No native text was detected and PDF OCR "
                                "fallback is disabled."
                            ],
                        )
                        blank_page_count += 1

                    extracted_character_count = (
                        self._checked_character_total(
                            upload,
                            current=extracted_character_count,
                            addition=len(page.text),
                        )
                    )
                    pages.append(page)
                finally:
                    if pdf_page is not None:
                        pdf_page.close()

        except DocumentParserError:
            raise
        except Exception as error:
            raise MalformedDocumentError(
                f"PDF '{upload.filename}' is malformed, encrypted, "
                "or unreadable."
            ) from error
        finally:
            if pdf_document is not None:
                pdf_document.close()

        for page in pages:
            warnings.extend(page.warnings)

        full_text = "\n\n".join(
            page.text for page in pages if page.text
        )

        if not full_text:
            raise EmptyDocumentError(
                f"PDF '{upload.filename}' contains no readable text."
            )

        self._checked_character_total(
            upload,
            current=0,
            addition=len(full_text),
        )
        blocks = [
            block
            for page in pages
            for block in page.blocks
        ]
        metadata_title = _normalise_inline_text(
            str(metadata.get("Title", ""))
        )
        title = metadata_title or _derive_title_from_blocks(
            blocks,
            fallback_filename=upload.filename,
        )

        return ParsedDocument(
            document_id=upload.document_id,
            title=title[:1000],
            pages=pages,
            full_text=full_text,
            page_count=len(pages),
            character_count=len(full_text),
            word_count=len(full_text.split()),
            parser_name=self.parser_name,
            parser_version=self.parser_version,
            extraction_method=_document_extraction_method(pages),
            extraction_confidence=_document_confidence(pages),
            warnings=warnings,
            errors=[],
            parser_metadata={
                "source_format": DocumentFormat.PDF.value,
                "effective_document_format": DocumentFormat.PDF.value,
                "native_page_count": native_page_count,
                "ocr_page_count": ocr_page_count,
                "blank_page_count": blank_page_count,
                "pdf_ocr_scale": self.config.pdf_ocr_scale,
                "pdf_ocr_total_pixels": total_ocr_pixels,
                "attachment_count": attachment_count,
                "document_metadata": metadata,
                "block_count": len(blocks),
            },
        )

    @staticmethod
    def _validated_pdf_page_size(
        pdf_page: Any,
        *,
        page_number: int,
    ) -> tuple[float, float]:
        width, height = pdf_page.get_size()
        width = float(width)
        height = float(height)

        if (
            not math.isfinite(width)
            or not math.isfinite(height)
            or width <= 0
            or height <= 0
        ):
            raise MalformedDocumentError(
                f"PDF page {page_number} has invalid dimensions."
            )

        return width, height

    @staticmethod
    def _extract_pdf_page_text(
        pdf_page: Any,
        *,
        page_number: int,
        warnings: list[str],
    ) -> str:
        text_page: Any = None

        try:
            text_page = pdf_page.get_textpage()
            return _normalise_document_text(
                text_page.get_text_range()
            )
        except Exception:
            warnings.append(
                f"Native PDF text extraction failed on page "
                f"{page_number}; OCR fallback was attempted."
            )
            return ""
        finally:
            if text_page is not None:
                text_page.close()

    @staticmethod
    def _build_native_pdf_page(
        *,
        text: str,
        page_number: int,
        width: float,
        height: float,
        sequence_number: int,
    ) -> tuple[ParsedPage, int]:
        source_blocks = _extract_text_blocks(
            text,
            markdown=False,
        )
        blocks: list[ParsedContentBlock] = []

        for source_block in source_blocks:
            attributes = dict(source_block.attributes)
            attributes["pdf_page_number"] = page_number
            blocks.append(
                source_block.model_copy(
                    update={
                        "page_number": page_number,
                        "sequence_number": sequence_number,
                        "attributes": attributes,
                    }
                )
            )
            sequence_number += 1

        return (
            ParsedPage(
                page_number=page_number,
                width=width,
                height=height,
                text=text,
                blocks=blocks,
                extraction_method=ExtractionMethod.NATIVE_TEXT,
                extraction_confidence=1.0,
                warnings=[],
            ),
            sequence_number,
        )

    def _build_ocr_pdf_page(
        self,
        *,
        upload: DocumentUpload,
        pdf_page: Any,
        page_number: int,
        width: float,
        height: float,
        sequence_number: int,
        current_ocr_pixels: int,
    ) -> tuple[ParsedPage, int, int]:
        scale = self.config.pdf_ocr_scale
        rendered_width = max(1, math.ceil(width * scale))
        rendered_height = max(1, math.ceil(height * scale))
        rendered_pixels = rendered_width * rendered_height

        if rendered_pixels > self.config.maximum_pdf_ocr_page_pixels:
            raise DocumentTooLargeError(
                f"Rendered PDF page {page_number} would contain "
                f"{rendered_pixels} pixels, which exceeds the "
                "configured per-page OCR limit of "
                f"{self.config.maximum_pdf_ocr_page_pixels}."
            )

        if (
            current_ocr_pixels + rendered_pixels
            > self.config.maximum_pdf_ocr_total_pixels
        ):
            raise DocumentTooLargeError(
                f"PDF OCR would exceed the configured total limit of "
                f"{self.config.maximum_pdf_ocr_total_pixels} pixels."
            )

        bitmap: Any = None

        try:
            bitmap = pdf_page.render(
                scale=scale,
                rotation=0,
                grayscale=True,
            )
            image = bitmap.to_pil().copy()
        except Exception as error:
            raise MalformedDocumentError(
                f"PDF page {page_number} could not be rendered for OCR."
            ) from error
        finally:
            if bitmap is not None:
                bitmap.close()

        image_buffer = BytesIO()
        image.save(image_buffer, format="PNG")
        image.close()
        image_content = image_buffer.getvalue()
        image_filename = (
            f"{Path(upload.filename).stem[:440]}-page-"
            f"{page_number}.png"
        )
        image_upload = upload.model_copy(
            update={
                "filename": image_filename,
                "original_filename": image_filename,
                "document_format": DocumentFormat.PNG,
                "media_type": "image/png",
                "size_bytes": len(image_content),
                "storage_key": (
                    f"{upload.storage_key[:980]}"
                    f"#pdf-page-{page_number}"
                ),
                "checksum_sha256": (
                    DocumentUpload.calculate_sha256(image_content)
                ),
                "password_protected": False,
            }
        )

        try:
            parsed_image = self._fallback_parser.parse(
                image_upload,
                image_content,
            )
        except EmptyDocumentError:
            return (
                ParsedPage(
                    page_number=page_number,
                    width=width,
                    height=height,
                    text="",
                    blocks=[],
                    extraction_method=ExtractionMethod.OCR,
                    extraction_confidence=0.0,
                    warnings=[
                        f"No readable native or OCR text was detected "
                        f"on PDF page {page_number}."
                    ],
                ),
                sequence_number,
                rendered_pixels,
            )
        except DocumentParserError as error:
            raise DocumentParserError(
                f"OCR failed for PDF page {page_number}: {error}"
            ) from error

        source_page = parsed_image.pages[0]
        blocks: list[ParsedContentBlock] = []

        for source_block in source_page.blocks:
            bounding_box = source_block.bounding_box

            if bounding_box is not None:
                bounding_box = BoundingBox(
                    x=bounding_box.x / scale,
                    y=bounding_box.y / scale,
                    width=bounding_box.width / scale,
                    height=bounding_box.height / scale,
                    page_width=width,
                    page_height=height,
                )

            attributes = dict(source_block.attributes)
            attributes["pdf_page_number"] = page_number
            blocks.append(
                source_block.model_copy(
                    update={
                        "page_number": page_number,
                        "sequence_number": sequence_number,
                        "bounding_box": bounding_box,
                        "attributes": attributes,
                    }
                )
            )
            sequence_number += 1

        page_warnings = [
            warning.replace(
                "page 1",
                f"page {page_number}",
            )
            for warning in source_page.warnings
        ]

        return (
            ParsedPage(
                page_number=page_number,
                width=width,
                height=height,
                text=source_page.text,
                blocks=blocks,
                extraction_method=ExtractionMethod.OCR,
                extraction_confidence=(
                    source_page.extraction_confidence
                ),
                warnings=page_warnings,
            ),
            sequence_number,
            rendered_pixels,
        )

    def _parse_docx(
        self,
        upload: DocumentUpload,
        content: bytes,
    ) -> ParsedDocument:
        archive_metadata = self._preflight_ooxml(
            upload,
            content,
            effective_format=DocumentFormat.DOCX,
        )

        try:
            document = load_docx_document(BytesIO(content))
            blocks: list[ParsedContentBlock] = []
            section_path: list[str] = []
            sequence_number = 1
            table_count = 0
            extracted_character_count = 0

            for item in document.iter_inner_content():
                if isinstance(item, DocxParagraph):
                    block = self._docx_paragraph_block(
                        item,
                        page_number=1,
                        sequence_number=sequence_number,
                        section_path=section_path,
                    )

                    if block is None:
                        continue

                    blocks.append(block)
                    sequence_number += 1
                    extracted_character_count = (
                        self._checked_character_total(
                            upload,
                            current=extracted_character_count,
                            addition=len(block.text),
                        )
                    )
                elif isinstance(item, DocxTable):
                    table = self._docx_table(
                        item,
                        upload=upload,
                    )

                    if table is None:
                        continue

                    table_text = _table_to_text(
                        table.headers,
                        table.rows,
                    )
                    blocks.append(
                        ParsedContentBlock(
                            block_type=ContentBlockType.TABLE,
                            text=table_text,
                            page_number=1,
                            section_path=list(section_path),
                            sequence_number=sequence_number,
                            table=table,
                            extraction_method=(
                                ExtractionMethod.TABLE_EXTRACTION
                            ),
                            extraction_confidence=1.0,
                            attributes={
                                "table_number": table_count + 1,
                            },
                        )
                    )
                    sequence_number += 1
                    table_count += 1
                    extracted_character_count = (
                        self._checked_character_total(
                            upload,
                            current=extracted_character_count,
                            addition=len(table_text),
                        )
                    )

                if len(blocks) > self.config.maximum_docx_blocks:
                    raise DocumentTooLargeError(
                        f"DOCX '{upload.filename}' contains more than "
                        f"{self.config.maximum_docx_blocks} content "
                        "blocks."
                    )

            if not blocks:
                raise EmptyDocumentError(
                    f"DOCX '{upload.filename}' contains no readable "
                    "text or tables."
                )

            page_text = "\n\n".join(
                block.text for block in blocks if block.text
            )
            self._checked_character_total(
                upload,
                current=0,
                addition=len(page_text),
            )
            properties = _docx_core_properties(document.core_properties)
            property_title = _normalise_inline_text(
                str(properties.get("title", ""))
            )
            title = property_title or _derive_title_from_blocks(
                blocks,
                fallback_filename=upload.filename,
            )
            embedded_image_count = len(document.inline_shapes)
            warnings = self._size_mismatch_warnings(upload, content)

            if embedded_image_count:
                warnings.append(
                    f"{embedded_image_count} embedded DOCX image(s) "
                    "were retained as metadata but were not "
                    "OCR-processed."
                )

            page_method = _block_extraction_method(blocks)
            page = ParsedPage(
                page_number=1,
                text=page_text,
                blocks=blocks,
                extraction_method=page_method,
                extraction_confidence=1.0,
                warnings=[],
            )

            return ParsedDocument(
                document_id=upload.document_id,
                title=title[:1000],
                pages=[page],
                full_text=page_text,
                page_count=1,
                character_count=len(page_text),
                word_count=len(page_text.split()),
                parser_name=self.parser_name,
                parser_version=self.parser_version,
                extraction_method=page_method,
                extraction_confidence=1.0,
                warnings=warnings,
                errors=[],
                parser_metadata={
                    "source_format": DocumentFormat.DOCX.value,
                    "effective_document_format": (
                        DocumentFormat.DOCX.value
                    ),
                    "block_count": len(blocks),
                    "table_count": table_count,
                    "section_count": len(document.sections),
                    "embedded_image_count": embedded_image_count,
                    "document_properties": properties,
                    **archive_metadata,
                },
            )
        except DocumentParserError:
            raise
        except Exception as error:
            raise MalformedDocumentError(
                f"DOCX '{upload.filename}' is malformed or unreadable."
            ) from error

    def _docx_paragraph_block(
        self,
        paragraph: DocxParagraph,
        *,
        page_number: int,
        sequence_number: int,
        section_path: list[str],
    ) -> ParsedContentBlock | None:
        text = _normalise_inline_text(
            paragraph.text.replace("\x00", "")
        )

        if not text:
            return None

        style_name = ""

        try:
            if paragraph.style is not None:
                style_name = paragraph.style.name or ""
        except (AttributeError, KeyError):
            style_name = ""

        style_name = _normalise_inline_text(style_name)
        style_key = style_name.casefold()
        block_type = ContentBlockType.PARAGRAPH
        attributes: dict[str, Any] = {}
        block_text = text
        heading_match = _HEADING_STYLE_PATTERN.fullmatch(style_name)

        if style_name:
            attributes["style_name"] = style_name

        if style_key == "title":
            block_type = ContentBlockType.TITLE
            section_path[:] = [text]
            attributes["heading_level"] = 1
        elif heading_match:
            heading_level = int(heading_match.group(1))
            block_type = ContentBlockType.HEADING
            section_path[:] = section_path[: heading_level - 1]
            section_path.append(text)
            attributes["heading_level"] = heading_level
        elif style_key.startswith("list"):
            block_type = ContentBlockType.LIST
            attributes["item_count"] = 1
        else:
            safety_match = _SAFETY_PREFIX_PATTERN.match(text)

            if safety_match:
                safety_label = safety_match.group(1).lower()
                block_text = safety_match.group(2)
                block_type = {
                    "danger": ContentBlockType.DANGER,
                    "warning": ContentBlockType.WARNING,
                    "caution": ContentBlockType.CAUTION,
                    "note": ContentBlockType.NOTE,
                }[safety_label]
                attributes["safety_label"] = safety_label

        return ParsedContentBlock(
            block_type=block_type,
            text=block_text,
            page_number=page_number,
            section_path=list(section_path),
            sequence_number=sequence_number,
            extraction_method=ExtractionMethod.NATIVE_TEXT,
            extraction_confidence=1.0,
            attributes=attributes,
        )

    def _docx_table(
        self,
        source_table: DocxTable,
        *,
        upload: DocumentUpload,
    ) -> ParsedTable | None:
        row_count = len(source_table.rows)

        if row_count > self.config.maximum_table_rows:
            raise DocumentTooLargeError(
                f"DOCX table in '{upload.filename}' contains "
                f"{row_count} rows, which exceeds the configured "
                f"limit of {self.config.maximum_table_rows}."
            )

        rows: list[list[str]] = []
        maximum_width = 0

        for source_row in source_table.rows:
            row = [
                self._format_cell_value(
                    cell.text,
                    upload=upload,
                )
                for cell in source_row.cells
            ]
            maximum_width = max(maximum_width, len(row))

            if any(row):
                rows.append(row)

        if maximum_width > self.config.maximum_table_columns:
            raise DocumentTooLargeError(
                f"DOCX table in '{upload.filename}' contains "
                f"{maximum_width} columns, which exceeds the "
                "configured limit of "
                f"{self.config.maximum_table_columns}."
            )

        if not rows:
            return None

        for row in rows:
            row.extend([""] * (maximum_width - len(row)))

        return ParsedTable(
            headers=rows[0],
            rows=rows[1:],
            column_count=maximum_width,
            row_count=max(0, len(rows) - 1),
        )

    def _parse_xlsx(
        self,
        upload: DocumentUpload,
        content: bytes,
    ) -> ParsedDocument:
        archive_metadata = self._preflight_ooxml(
            upload,
            content,
            effective_format=DocumentFormat.XLSX,
        )
        workbook: Any = None

        try:
            workbook = load_workbook(
                filename=BytesIO(content),
                read_only=True,
                data_only=True,
                keep_links=False,
            )

            if len(workbook.worksheets) > (
                self.config.maximum_workbook_sheets
            ):
                raise DocumentTooLargeError(
                    f"XLSX '{upload.filename}' contains "
                    f"{len(workbook.worksheets)} sheets, which exceeds "
                    "the configured limit of "
                    f"{self.config.maximum_workbook_sheets}."
                )

            pages: list[ParsedPage] = []
            sequence_number = 1
            scanned_cell_count = 0
            extracted_character_count = 0

            for page_number, worksheet in enumerate(
                workbook.worksheets,
                start=1,
            ):
                maximum_row = int(worksheet.max_row or 0)
                maximum_column = int(worksheet.max_column or 0)
                scanned_cell_count = self._checked_sheet_dimensions(
                    upload,
                    sheet_name=worksheet.title,
                    row_count=maximum_row,
                    column_count=maximum_column,
                    current_cell_count=scanned_cell_count,
                )
                row_records: list[tuple[int, list[str]]] = []

                if maximum_row and maximum_column:
                    for row_number, source_row in enumerate(
                        worksheet.iter_rows(
                            min_row=1,
                            max_row=maximum_row,
                            min_col=1,
                            max_col=maximum_column,
                            values_only=True,
                        ),
                        start=1,
                    ):
                        row = [
                            self._format_cell_value(
                                value,
                                upload=upload,
                            )
                            for value in source_row
                        ]

                        while row and not row[-1]:
                            row.pop()

                        if any(row):
                            row_records.append((row_number, row))

                (
                    page,
                    sequence_number,
                    page_character_count,
                ) = self._build_spreadsheet_page(
                    sheet_name=worksheet.title,
                    page_number=page_number,
                    row_records=row_records,
                    sequence_number=sequence_number,
                    source_format=DocumentFormat.XLSX,
                    sheet_state=worksheet.sheet_state,
                )
                extracted_character_count = (
                    self._checked_character_total(
                        upload,
                        current=extracted_character_count,
                        addition=page_character_count,
                    )
                )
                pages.append(page)

            return self._build_spreadsheet_document(
                upload=upload,
                content=content,
                pages=pages,
                source_format=DocumentFormat.XLSX,
                workbook_title=_normalise_inline_text(
                    str(workbook.properties.title or "")
                ),
                workbook_metadata=_openpyxl_properties(
                    workbook.properties
                ),
                scanned_cell_count=scanned_cell_count,
                archive_metadata=archive_metadata,
            )
        except DocumentParserError:
            raise
        except Exception as error:
            raise MalformedDocumentError(
                f"XLSX '{upload.filename}' is malformed or unreadable."
            ) from error
        finally:
            if workbook is not None:
                workbook.close()

    def _parse_xls(
        self,
        upload: DocumentUpload,
        content: bytes,
    ) -> ParsedDocument:
        if not content.startswith(_OLE_COMPOUND_FILE_SIGNATURE):
            raise MalformedDocumentError(
                f"XLS '{upload.filename}' does not contain a valid "
                "OLE compound-file header."
            )

        workbook: Any = None

        try:
            workbook = xlrd.open_workbook(
                file_contents=content,
                on_demand=True,
            )

            if workbook.nsheets > self.config.maximum_workbook_sheets:
                raise DocumentTooLargeError(
                    f"XLS '{upload.filename}' contains "
                    f"{workbook.nsheets} sheets, which exceeds the "
                    "configured limit of "
                    f"{self.config.maximum_workbook_sheets}."
                )

            pages: list[ParsedPage] = []
            sequence_number = 1
            scanned_cell_count = 0
            extracted_character_count = 0

            for page_number in range(1, workbook.nsheets + 1):
                sheet = workbook.sheet_by_index(page_number - 1)
                scanned_cell_count = self._checked_sheet_dimensions(
                    upload,
                    sheet_name=sheet.name,
                    row_count=sheet.nrows,
                    column_count=sheet.ncols,
                    current_cell_count=scanned_cell_count,
                )
                row_records: list[tuple[int, list[str]]] = []

                for row_index in range(sheet.nrows):
                    row = [
                        self._format_xls_cell(
                            sheet.cell(row_index, column_index),
                            workbook=workbook,
                            upload=upload,
                        )
                        for column_index in range(sheet.ncols)
                    ]

                    while row and not row[-1]:
                        row.pop()

                    if any(row):
                        row_records.append((row_index + 1, row))

                (
                    page,
                    sequence_number,
                    page_character_count,
                ) = self._build_spreadsheet_page(
                    sheet_name=sheet.name,
                    page_number=page_number,
                    row_records=row_records,
                    sequence_number=sequence_number,
                    source_format=DocumentFormat.XLS,
                    sheet_state="visible",
                )
                extracted_character_count = (
                    self._checked_character_total(
                        upload,
                        current=extracted_character_count,
                        addition=page_character_count,
                    )
                )
                pages.append(page)

            return self._build_spreadsheet_document(
                upload=upload,
                content=content,
                pages=pages,
                source_format=DocumentFormat.XLS,
                workbook_title="",
                workbook_metadata={
                    "date_mode": workbook.datemode,
                    "sheet_names": workbook.sheet_names(),
                },
                scanned_cell_count=scanned_cell_count,
                archive_metadata={},
            )
        except DocumentParserError:
            raise
        except Exception as error:
            raise MalformedDocumentError(
                f"XLS '{upload.filename}' is malformed, encrypted, "
                "or unreadable."
            ) from error
        finally:
            if workbook is not None:
                workbook.release_resources()

    def _checked_sheet_dimensions(
        self,
        upload: DocumentUpload,
        *,
        sheet_name: str,
        row_count: int,
        column_count: int,
        current_cell_count: int,
    ) -> int:
        if row_count > self.config.maximum_spreadsheet_rows:
            raise DocumentTooLargeError(
                f"Sheet '{sheet_name}' in '{upload.filename}' contains "
                f"{row_count} rows, which exceeds the configured "
                f"limit of {self.config.maximum_spreadsheet_rows}."
            )

        if column_count > self.config.maximum_spreadsheet_columns:
            raise DocumentTooLargeError(
                f"Sheet '{sheet_name}' in '{upload.filename}' contains "
                f"{column_count} columns, which exceeds the configured "
                f"limit of {self.config.maximum_spreadsheet_columns}."
            )

        sheet_cell_count = row_count * column_count
        total_cell_count = current_cell_count + sheet_cell_count

        if total_cell_count > self.config.maximum_spreadsheet_cells:
            raise DocumentTooLargeError(
                f"Workbook '{upload.filename}' requires scanning "
                f"{total_cell_count} cells, which exceeds the "
                "configured limit of "
                f"{self.config.maximum_spreadsheet_cells}."
            )

        return total_cell_count

    def _build_spreadsheet_page(
        self,
        *,
        sheet_name: str,
        page_number: int,
        row_records: Sequence[tuple[int, list[str]]],
        sequence_number: int,
        source_format: DocumentFormat,
        sheet_state: str,
    ) -> tuple[ParsedPage, int, int]:
        if not row_records:
            return (
                ParsedPage(
                    page_number=page_number,
                    text="",
                    blocks=[],
                    extraction_method=(
                        ExtractionMethod.TABLE_EXTRACTION
                    ),
                    extraction_confidence=1.0,
                    warnings=[
                        f"Worksheet '{sheet_name}' contains no "
                        "readable cell values."
                    ],
                ),
                sequence_number,
                0,
            )

        maximum_width = max(
            len(row) for _, row in row_records
        )
        padded_rows = []

        for _, row in row_records:
            padded_rows.append(
                [*row, *([""] * (maximum_width - len(row)))]
            )

        headers = padded_rows[0]
        rows = padded_rows[1:]
        table = ParsedTable(
            headers=headers,
            rows=rows,
            caption=sheet_name,
            column_count=maximum_width,
            row_count=len(rows),
        )
        table_text = _table_to_text(headers, rows)
        first_row_number = row_records[0][0]
        last_row_number = row_records[-1][0]
        end_column = get_column_letter(maximum_width)
        block = ParsedContentBlock(
            block_type=ContentBlockType.TABLE,
            text=table_text,
            page_number=page_number,
            sequence_number=sequence_number,
            spreadsheet_range=SpreadsheetCellRange(
                sheet_name=sheet_name,
                start_cell=f"A{first_row_number}",
                end_cell=f"{end_column}{last_row_number}",
            ),
            table=table,
            extraction_method=ExtractionMethod.TABLE_EXTRACTION,
            extraction_confidence=1.0,
            attributes={
                "source_format": source_format.value,
                "sheet_name": sheet_name,
                "sheet_state": sheet_state,
            },
        )

        return (
            ParsedPage(
                page_number=page_number,
                text=table_text,
                blocks=[block],
                extraction_method=ExtractionMethod.TABLE_EXTRACTION,
                extraction_confidence=1.0,
                warnings=[],
            ),
            sequence_number + 1,
            len(table_text),
        )

    def _build_spreadsheet_document(
        self,
        *,
        upload: DocumentUpload,
        content: bytes,
        pages: list[ParsedPage],
        source_format: DocumentFormat,
        workbook_title: str,
        workbook_metadata: dict[str, Any],
        scanned_cell_count: int,
        archive_metadata: dict[str, Any],
    ) -> ParsedDocument:
        full_text = "\n\n".join(
            page.text for page in pages if page.text
        )

        if not full_text:
            raise EmptyDocumentError(
                f"Workbook '{upload.filename}' contains no readable "
                "cell values."
            )

        self._checked_character_total(
            upload,
            current=0,
            addition=len(full_text),
        )
        warnings = self._size_mismatch_warnings(upload, content)
        for page in pages:
            warnings.extend(page.warnings)

        return ParsedDocument(
            document_id=upload.document_id,
            title=(
                workbook_title
                or Path(upload.filename).stem
                or upload.filename
            )[:1000],
            pages=pages,
            full_text=full_text,
            page_count=len(pages),
            character_count=len(full_text),
            word_count=len(full_text.split()),
            parser_name=self.parser_name,
            parser_version=self.parser_version,
            extraction_method=ExtractionMethod.TABLE_EXTRACTION,
            extraction_confidence=1.0,
            warnings=warnings,
            errors=[],
            parser_metadata={
                "source_format": source_format.value,
                "effective_document_format": source_format.value,
                "sheet_count": len(pages),
                "non_empty_sheet_count": sum(
                    1 for page in pages if page.text
                ),
                "scanned_cell_count": scanned_cell_count,
                "formula_values": "cached",
                "workbook_metadata": workbook_metadata,
                **archive_metadata,
            },
        )

    def _format_xls_cell(
        self,
        cell: Any,
        *,
        workbook: Any,
        upload: DocumentUpload,
    ) -> str:
        if cell.ctype in {
            xlrd.XL_CELL_EMPTY,
            xlrd.XL_CELL_BLANK,
        }:
            return ""

        if cell.ctype == xlrd.XL_CELL_DATE:
            try:
                value: Any = xlrd.xldate_as_datetime(
                    cell.value,
                    workbook.datemode,
                )
            except (OverflowError, TypeError, ValueError):
                value = cell.value
        elif cell.ctype == xlrd.XL_CELL_BOOLEAN:
            value = bool(cell.value)
        elif cell.ctype == xlrd.XL_CELL_ERROR:
            value = error_text_from_code.get(
                cell.value,
                f"#ERROR({cell.value})",
            )
        else:
            value = cell.value

        return self._format_cell_value(value, upload=upload)

    def _format_cell_value(
        self,
        value: Any,
        *,
        upload: DocumentUpload,
    ) -> str:
        if value is None:
            text = ""
        elif isinstance(value, bool):
            text = "TRUE" if value else "FALSE"
        elif isinstance(value, datetime):
            text = value.isoformat(sep=" ")
        elif isinstance(value, (date, time)):
            text = value.isoformat()
        elif isinstance(value, timedelta):
            text = str(value)
        elif isinstance(value, Decimal):
            text = format(value, "f")
        elif isinstance(value, int):
            text = str(value)
        elif isinstance(value, float):
            if not math.isfinite(value):
                raise MalformedDocumentError(
                    f"Workbook '{upload.filename}' contains a "
                    "non-finite numeric cell value."
                )
            text = str(int(value)) if value.is_integer() else str(value)
        elif isinstance(value, bytes):
            text = value.decode("utf-8", errors="replace")
        else:
            text = str(value)

        text = _normalise_inline_text(text.replace("\x00", ""))

        if len(text) > self.config.maximum_cell_characters:
            raise DocumentTooLargeError(
                f"A cell in '{upload.filename}' contains "
                f"{len(text)} characters, which exceeds the "
                "configured limit of "
                f"{self.config.maximum_cell_characters}."
            )

        return text

    def _preflight_ooxml(
        self,
        upload: DocumentUpload,
        content: bytes,
        *,
        effective_format: DocumentFormat,
    ) -> dict[str, Any]:
        if not content.startswith(
            (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")
        ):
            raise MalformedDocumentError(
                f"{effective_format.value.upper()} "
                f"'{upload.filename}' does not contain a valid ZIP "
                "container header."
            )

        member_names: set[str] = set()
        total_uncompressed_bytes = 0
        xml_member_count = 0

        try:
            with ZipFile(BytesIO(content)) as archive:
                members = archive.infolist()

                if len(members) > self.config.maximum_archive_members:
                    raise DocumentTooLargeError(
                        f"OOXML archive '{upload.filename}' contains "
                        f"{len(members)} members, which exceeds the "
                        "configured limit of "
                        f"{self.config.maximum_archive_members}."
                    )

                for member in members:
                    member_name = _safe_archive_member_name(
                        member.filename,
                        upload=upload,
                    )

                    if member_name in member_names:
                        raise UnsafeOfficeArchiveError(
                            f"OOXML archive '{upload.filename}' contains "
                            f"a duplicate member '{member_name}'."
                        )

                    member_names.add(member_name)

                    if member.flag_bits & 0x1:
                        raise PasswordProtectedDocumentError(
                            f"OOXML archive '{upload.filename}' contains "
                            f"an encrypted member '{member_name}'."
                        )

                    if member_name.casefold().endswith(
                        "vbaproject.bin"
                    ):
                        raise UnsafeOfficeArchiveError(
                            f"OOXML archive '{upload.filename}' contains "
                            "VBA macros, which are not accepted."
                        )

                    total_uncompressed_bytes += member.file_size

                    if total_uncompressed_bytes > (
                        self.config.maximum_archive_uncompressed_bytes
                    ):
                        raise DocumentTooLargeError(
                            f"OOXML archive '{upload.filename}' expands "
                            "to more than {} bytes.".format(
                                self.config
                                .maximum_archive_uncompressed_bytes
                            )
                        )

                    if member.file_size:
                        compression_ratio = (
                            member.file_size
                            / max(1, member.compress_size)
                        )

                        if compression_ratio > (
                            self.config
                            .maximum_archive_compression_ratio
                        ):
                            raise UnsafeOfficeArchiveError(
                                f"OOXML member '{member_name}' has an "
                                f"unsafe compression ratio of "
                                f"{compression_ratio:.1f}."
                            )

                    if (
                        not member.is_dir()
                        and member_name.casefold().endswith(
                            _OOXML_XML_SUFFIXES
                        )
                    ):
                        xml_member_count += 1
                        xml_content = archive.read(member)
                        self._validate_ooxml_xml(
                            upload,
                            member_name=member_name,
                            content=xml_content,
                        )

                missing_members = (
                    _OOXML_REQUIRED_MEMBERS[effective_format]
                    - member_names
                )

                if missing_members:
                    raise MalformedDocumentError(
                        f"OOXML archive '{upload.filename}' is missing "
                        "required member(s): "
                        f"{', '.join(sorted(missing_members))}."
                    )
        except DocumentParserError:
            raise
        except BadZipFile as error:
            raise MalformedDocumentError(
                f"OOXML archive '{upload.filename}' is not a valid "
                "ZIP container."
            ) from error
        except (OSError, RuntimeError) as error:
            raise MalformedDocumentError(
                f"OOXML archive '{upload.filename}' could not be "
                "read safely."
            ) from error

        return {
            "archive_member_count": len(member_names),
            "archive_xml_member_count": xml_member_count,
            "archive_uncompressed_bytes": total_uncompressed_bytes,
        }

    @staticmethod
    def _validate_ooxml_xml(
        upload: DocumentUpload,
        *,
        member_name: str,
        content: bytes,
    ) -> None:
        try:
            SafeElementTree.fromstring(
                content,
                forbid_dtd=True,
                forbid_entities=True,
                forbid_external=True,
            )
        except (DefusedXmlException, ElementTree.ParseError) as error:
            raise UnsafeOfficeArchiveError(
                f"OOXML member '{member_name}' in "
                f"'{upload.filename}' contains unsafe or malformed XML."
            ) from error

    def _checked_character_total(
        self,
        upload: DocumentUpload,
        *,
        current: int,
        addition: int,
    ) -> int:
        total = current + addition

        if total > self.config.maximum_extracted_characters:
            raise DocumentTooLargeError(
                f"Document '{upload.filename}' contains more than "
                f"{self.config.maximum_extracted_characters} extracted "
                "characters."
            )

        return total

    @staticmethod
    def _size_mismatch_warnings(
        upload: DocumentUpload,
        content: bytes,
    ) -> list[str]:
        if upload.size_bytes == len(content):
            return []

        return [
            "Upload size metadata does not match the supplied document "
            f"content size: metadata={upload.size_bytes}, "
            f"actual={len(content)}."
        ]


def parse_pdf_office_document(
    upload: DocumentUpload,
    content: bytes,
    *,
    config: PdfOfficeDocumentParserConfig | None = None,
) -> ParsedDocument:
    """Convenience function for the complete parser chain."""

    return PdfOfficeDocumentParser(config=config).parse(upload, content)


def _safe_archive_member_name(
    member_name: str,
    *,
    upload: DocumentUpload,
) -> str:
    normalised = member_name.replace("\\", "/")
    path = PurePosixPath(normalised)

    if (
        not normalised
        or "\x00" in normalised
        or path.is_absolute()
        or ".." in path.parts
        or (
            path.parts
            and ":" in path.parts[0]
        )
    ):
        raise UnsafeOfficeArchiveError(
            f"OOXML archive '{upload.filename}' contains an unsafe "
            f"member path '{member_name}'."
        )

    return path.as_posix()


def _document_extraction_method(
    pages: Sequence[ParsedPage],
) -> ExtractionMethod:
    methods = {
        page.extraction_method
        for page in pages
        if page.text
    }

    if not methods:
        return ExtractionMethod.UNKNOWN

    if len(methods) == 1:
        return next(iter(methods))

    return ExtractionMethod.HYBRID


def _block_extraction_method(
    blocks: Sequence[ParsedContentBlock],
) -> ExtractionMethod:
    methods = {
        block.extraction_method
        for block in blocks
    }

    if not methods:
        return ExtractionMethod.UNKNOWN

    if len(methods) == 1:
        return next(iter(methods))

    return ExtractionMethod.HYBRID


def _document_confidence(
    pages: Sequence[ParsedPage],
) -> float:
    confidences = [
        page.extraction_confidence
        for page in pages
        if page.text
    ]

    if not confidences:
        return 0.0

    return sum(confidences) / len(confidences)


def _metadata_value(value: Any) -> Any:
    if isinstance(value, datetime):
        return value.isoformat(sep=" ")

    if isinstance(value, (date, time)):
        return value.isoformat()

    if isinstance(value, timedelta):
        return str(value)

    if isinstance(value, (str, int, float, bool)) or value is None:
        return value

    return str(value)


def _docx_core_properties(properties: Any) -> dict[str, Any]:
    field_names = (
        "author",
        "category",
        "comments",
        "content_status",
        "created",
        "identifier",
        "keywords",
        "language",
        "last_modified_by",
        "last_printed",
        "modified",
        "revision",
        "subject",
        "title",
        "version",
    )

    return {
        field_name: _metadata_value(
            getattr(properties, field_name, None)
        )
        for field_name in field_names
        if getattr(properties, field_name, None) not in (None, "")
    }


def _openpyxl_properties(properties: Any) -> dict[str, Any]:
    field_names = (
        "category",
        "contentStatus",
        "created",
        "creator",
        "description",
        "identifier",
        "keywords",
        "language",
        "lastModifiedBy",
        "lastPrinted",
        "modified",
        "revision",
        "subject",
        "title",
        "version",
    )

    return {
        field_name: _metadata_value(
            getattr(properties, field_name, None)
        )
        for field_name in field_names
        if getattr(properties, field_name, None) not in (None, "")
    }
