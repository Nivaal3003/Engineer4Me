"""Tests for bounded PDF and Office document parsing."""

from __future__ import annotations

from dataclasses import FrozenInstanceError, dataclass, replace
from datetime import datetime
from hashlib import sha256
from io import BytesIO
from typing import Any
from unittest import TestCase
from unittest.mock import patch
import warnings
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from openpyxl import Workbook
from PIL import Image
import xlrd

from app.ingestion.document_models import (
    BoundingBox,
    ContentBlockType,
    DocumentFormat,
    DocumentSource,
    DocumentUpload,
    ExtractionMethod,
    ParsedContentBlock,
    ParsedDocument,
    ParsedPage,
)
from app.ingestion.document_parser import (
    DocumentParserError,
    DocumentTooLargeError,
    EmptyDocumentError,
    MalformedDocumentError,
    UnsupportedDocumentFormatError,
)
from app.ingestion.pdf_office_document_parser import (
    PasswordProtectedDocumentError,
    PdfOfficeDocumentParser,
    PdfOfficeDocumentParserConfig,
    UnsafeOfficeArchiveError,
    parse_pdf_office_document,
)
from app.ingestion.ocr_document_parser import (
    OcrAwareDocumentParser,
    OcrAwareDocumentParserConfig,
)


_OLE_SIGNATURE = bytes.fromhex("D0CF11E0A1B11AE1")


def build_upload(
    *,
    filename: str,
    document_format: DocumentFormat,
    content: bytes,
    media_type: str | None = None,
    declared_size: int | None = None,
    password_protected: bool = False,
) -> DocumentUpload:
    """Build valid upload metadata for one parser test."""

    return DocumentUpload(
        filename=filename,
        document_format=document_format,
        media_type=media_type,
        size_bytes=(
            len(content)
            if declared_size is None
            else declared_size
        ),
        storage_key=f"uploads/{filename}",
        checksum_sha256=sha256(content).hexdigest(),
        source=DocumentSource(
            source_name="PDF and Office parser test suite",
        ),
        password_protected=password_protected,
    )


def _pdf_literal(value: str) -> str:
    """Escape one ASCII PDF literal string."""

    return (
        value.replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
    )


def build_pdf_bytes(
    *page_texts: str,
    title: str | None = None,
    width: int = 200,
    height: int = 100,
) -> bytes:
    """Build a deterministic, valid PDF without a test-only dependency."""

    if not page_texts:
        page_texts = ("",)

    page_count = len(page_texts)
    page_object_ids = list(range(3, 3 + page_count))
    content_object_ids = list(
        range(3 + page_count, 3 + (2 * page_count))
    )
    font_object_id = 3 + (2 * page_count)
    info_object_id = font_object_id + 1 if title else None
    maximum_object_id = info_object_id or font_object_id
    objects: dict[int, bytes] = {}
    kids = " ".join(
        f"{object_id} 0 R"
        for object_id in page_object_ids
    )

    objects[1] = b"<< /Type /Catalog /Pages 2 0 R >>"
    objects[2] = (
        f"<< /Type /Pages /Kids [{kids}] "
        f"/Count {page_count} >>"
    ).encode("ascii")
    objects[font_object_id] = (
        b"<< /Type /Font /Subtype /Type1 "
        b"/BaseFont /Helvetica >>"
    )

    for index, page_text in enumerate(page_texts):
        page_object_id = page_object_ids[index]
        content_object_id = content_object_ids[index]
        objects[page_object_id] = (
            f"<< /Type /Page /Parent 2 0 R "
            f"/MediaBox [0 0 {width} {height}] "
            f"/Resources << /Font << "
            f"/F1 {font_object_id} 0 R >> >> "
            f"/Contents {content_object_id} 0 R >>"
        ).encode("ascii")

        if page_text:
            escaped_text = _pdf_literal(page_text)
            stream = (
                f"BT /F1 12 Tf 12 {height - 24} Td "
                f"({escaped_text}) Tj ET"
            ).encode("ascii")
        else:
            stream = b""

        objects[content_object_id] = (
            f"<< /Length {len(stream)} >>\nstream\n".encode(
                "ascii"
            )
            + stream
            + b"\nendstream"
        )

    if info_object_id is not None:
        objects[info_object_id] = (
            f"<< /Title ({_pdf_literal(title or '')}) >>"
        ).encode("ascii")

    output = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets: dict[int, int] = {}

    for object_id in range(1, maximum_object_id + 1):
        offsets[object_id] = len(output)
        output.extend(
            f"{object_id} 0 obj\n".encode("ascii")
        )
        output.extend(objects[object_id])
        output.extend(b"\nendobj\n")

    xref_offset = len(output)
    output.extend(
        f"xref\n0 {maximum_object_id + 1}\n".encode("ascii")
    )
    output.extend(b"0000000000 65535 f \n")

    for object_id in range(1, maximum_object_id + 1):
        output.extend(
            f"{offsets[object_id]:010d} 00000 n \n".encode(
                "ascii"
            )
        )

    trailer_parts = [
        f"/Size {maximum_object_id + 1}",
        "/Root 1 0 R",
    ]

    if info_object_id is not None:
        trailer_parts.append(f"/Info {info_object_id} 0 R")

    output.extend(
        (
            f"trailer\n<< {' '.join(trailer_parts)} >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(output)


def build_docx_bytes(
    *,
    include_content: bool = True,
    include_table: bool = True,
    include_image: bool = False,
) -> bytes:
    """Build a small deterministic DOCX test document."""

    document = Document()
    document.core_properties.title = "Valve Maintenance Manual"
    document.core_properties.author = "Engineer4Me"

    if include_content:
        document.add_heading(
            "Valve Maintenance Manual",
            level=0,
        )
        document.add_heading("Safety", level=1)
        document.add_paragraph(
            "DANGER: Depressurise the valve before service."
        )
        document.add_paragraph(
            "Inspect the actuator linkage.",
            style="List Bullet",
        )

    if include_table:
        table = document.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Model"
        table.cell(0, 1).text = "Maximum pressure"
        table.cell(1, 0).text = "CV-100"
        table.cell(1, 1).text = "16 bar"
        table.cell(2, 0).text = "CV-200"
        table.cell(2, 1).text = "25 bar"

    if include_image:
        image_buffer = BytesIO()
        image = Image.new(
            "RGB",
            (10, 10),
            color="white",
        )
        image.save(image_buffer, format="PNG")
        image.close()
        image_buffer.seek(0)
        document.add_picture(image_buffer)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_xlsx_bytes(
    *,
    include_values: bool = True,
    include_blank_sheet: bool = False,
    long_value: str | None = None,
) -> bytes:
    """Build a small deterministic XLSX workbook."""

    workbook = Workbook()
    workbook.properties.title = "Instrument Catalogue"
    workbook.properties.creator = "Engineer4Me"
    worksheet = workbook.active
    worksheet.title = "Specifications"

    if include_values:
        worksheet.append(
            [
                "Model",
                "Range",
                "Enabled",
                "Calibration date",
            ]
        )
        worksheet.append(
            [
                long_value or "PT-100",
                100,
                True,
                datetime(2026, 1, 2, 3, 4, 5),
            ]
        )

    if include_blank_sheet:
        blank_sheet = workbook.create_sheet("Archived")
        blank_sheet.sheet_state = "hidden"

    output = BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def append_zip_member(
    content: bytes,
    *,
    member_name: str,
    member_content: bytes,
) -> bytes:
    """Append one compressed member to an OOXML archive."""

    output = BytesIO(content)

    with ZipFile(
        output,
        mode="a",
        compression=ZIP_DEFLATED,
    ) as archive:
        archive.writestr(member_name, member_content)

    return output.getvalue()


def build_ocr_result(
    *,
    document_id: Any,
    text: str = "OCR PRESSURE TRANSMITTER",
    confidence: float = 0.86,
) -> ParsedDocument:
    """Build the result returned by a mocked image OCR parser."""

    block = ParsedContentBlock(
        block_type=ContentBlockType.PARAGRAPH,
        text=text,
        page_number=1,
        sequence_number=0,
        bounding_box=BoundingBox(
            x=0,
            y=0,
            width=400,
            height=200,
            page_width=400,
            page_height=200,
        ),
        extraction_method=ExtractionMethod.OCR,
        extraction_confidence=confidence,
        attributes={
            "ocr_engine": "tesseract",
        },
    )
    page = ParsedPage(
        page_number=1,
        width=400,
        height=200,
        text=text,
        blocks=[block],
        extraction_method=ExtractionMethod.OCR,
        extraction_confidence=confidence,
    )

    return ParsedDocument(
        document_id=document_id,
        pages=[page],
        parser_name="mock-image-ocr-parser",
        parser_version="1.0.0",
        extraction_method=ExtractionMethod.OCR,
        extraction_confidence=confidence,
    )


@dataclass
class FakeXlsCell:
    """Minimal xlrd cell used by legacy-workbook unit tests."""

    ctype: int
    value: Any


class FakeXlsSheet:
    """Minimal xlrd sheet used by legacy-workbook unit tests."""

    def __init__(
        self,
        name: str,
        rows: list[list[FakeXlsCell]],
    ) -> None:
        self.name = name
        self._rows = rows
        self.nrows = len(rows)
        self.ncols = max(
            (len(row) for row in rows),
            default=0,
        )

    def cell(
        self,
        row_index: int,
        column_index: int,
    ) -> FakeXlsCell:
        row = self._rows[row_index]

        if column_index >= len(row):
            return FakeXlsCell(xlrd.XL_CELL_EMPTY, "")

        return row[column_index]


class FakeXlsWorkbook:
    """Minimal xlrd workbook used by legacy-workbook unit tests."""

    def __init__(
        self,
        sheets: list[FakeXlsSheet],
    ) -> None:
        self._sheets = sheets
        self.nsheets = len(sheets)
        self.datemode = 0
        self.resources_released = False

    def sheet_by_index(self, index: int) -> FakeXlsSheet:
        return self._sheets[index]

    def sheet_names(self) -> list[str]:
        return [sheet.name for sheet in self._sheets]

    def release_resources(self) -> None:
        self.resources_released = True


def build_fake_xls_workbook() -> FakeXlsWorkbook:
    """Build a legacy workbook containing common xlrd cell types."""

    return FakeXlsWorkbook(
        [
            FakeXlsSheet(
                "Legacy equipment",
                [
                    [
                        FakeXlsCell(
                            xlrd.XL_CELL_TEXT,
                            "Model",
                        ),
                        FakeXlsCell(
                            xlrd.XL_CELL_TEXT,
                            "Range",
                        ),
                        FakeXlsCell(
                            xlrd.XL_CELL_TEXT,
                            "Enabled",
                        ),
                        FakeXlsCell(
                            xlrd.XL_CELL_TEXT,
                            "Commissioned",
                        ),
                        FakeXlsCell(
                            xlrd.XL_CELL_TEXT,
                            "Status",
                        ),
                    ],
                    [
                        FakeXlsCell(
                            xlrd.XL_CELL_TEXT,
                            "LT-100",
                        ),
                        FakeXlsCell(
                            xlrd.XL_CELL_NUMBER,
                            25.0,
                        ),
                        FakeXlsCell(
                            xlrd.XL_CELL_BOOLEAN,
                            1,
                        ),
                        FakeXlsCell(
                            xlrd.XL_CELL_DATE,
                            1.0,
                        ),
                        FakeXlsCell(
                            xlrd.XL_CELL_ERROR,
                            0x07,
                        ),
                    ],
                ],
            )
        ]
    )


class PdfOfficeParserConfigurationTests(TestCase):
    """Validate immutable and bounded parser configuration."""

    def test_defaults_are_bounded_and_frozen(self) -> None:
        config = PdfOfficeDocumentParserConfig()

        self.assertIsInstance(
            config.fallback_parser,
            OcrAwareDocumentParserConfig,
        )
        self.assertEqual(config.maximum_pdf_pages, 500)
        self.assertTrue(config.enable_pdf_ocr_fallback)
        self.assertEqual(config.pdf_ocr_scale, 2.0)
        self.assertEqual(
            config.maximum_spreadsheet_columns,
            16_384,
        )

        with self.assertRaises(FrozenInstanceError):
            config.maximum_pdf_pages = 1  # type: ignore[misc]

    def test_invalid_positive_controls_are_rejected(self) -> None:
        config = PdfOfficeDocumentParserConfig()

        for field_name in (
            "maximum_document_bytes",
            "maximum_extracted_characters",
            "maximum_pdf_pages",
            "minimum_pdf_native_characters",
            "maximum_archive_members",
            "maximum_docx_blocks",
            "maximum_spreadsheet_cells",
            "maximum_cell_characters",
        ):
            with self.subTest(field_name=field_name):
                with self.assertRaises(ValueError):
                    replace(config, **{field_name: 0})

    def test_invalid_special_controls_are_rejected(self) -> None:
        with self.assertRaises(TypeError):
            PdfOfficeDocumentParserConfig(
                fallback_parser=object(),  # type: ignore[arg-type]
            )

        with self.assertRaises(TypeError):
            PdfOfficeDocumentParserConfig(
                enable_pdf_ocr_fallback=1,  # type: ignore[arg-type]
            )

        with self.assertRaises(ValueError):
            PdfOfficeDocumentParserConfig(
                pdf_ocr_scale=float("inf"),
            )

        with self.assertRaises(ValueError):
            PdfOfficeDocumentParserConfig(
                maximum_pdf_ocr_page_pixels=100,
                maximum_pdf_ocr_total_pixels=99,
            )


class PdfOfficeParserRoutingTests(TestCase):
    """Validate format resolution and fallback parser delegation."""

    def test_parser_identity_and_fallback_are_exposed(self) -> None:
        parser = PdfOfficeDocumentParser()

        self.assertEqual(
            parser.parser_name,
            "engineer4me-pdf-office-document-parser",
        )
        self.assertEqual(parser.parser_version, "1.0.0")
        self.assertIsInstance(
            parser.fallback_parser,
            OcrAwareDocumentParser,
        )

    def test_supported_pdf_and_office_formats_are_reported(self) -> None:
        parser = PdfOfficeDocumentParser()

        for filename, document_format in (
            ("manual.pdf", DocumentFormat.PDF),
            ("manual.docx", DocumentFormat.DOCX),
            ("catalogue.xlsx", DocumentFormat.XLSX),
            ("legacy.xls", DocumentFormat.XLS),
        ):
            with self.subTest(document_format=document_format):
                upload = build_upload(
                    filename=filename,
                    document_format=document_format,
                    content=b"placeholder",
                )
                self.assertTrue(parser.supports(upload))

    def test_unknown_office_formats_are_resolved_from_suffixes(self) -> None:
        parser = PdfOfficeDocumentParser()

        for filename, expected_format in (
            ("manual.PDF", DocumentFormat.PDF),
            ("manual.DOCX", DocumentFormat.DOCX),
            ("catalogue.XLSX", DocumentFormat.XLSX),
            ("legacy.XLS", DocumentFormat.XLS),
        ):
            with self.subTest(filename=filename):
                upload = build_upload(
                    filename=filename,
                    document_format=DocumentFormat.UNKNOWN,
                    content=b"placeholder",
                )
                self.assertEqual(
                    parser.resolve_document_format(upload),
                    expected_format,
                )

    def test_text_document_is_delegated_to_standard_chain(self) -> None:
        content = (
            b"# Pressure Transmitter\n\n"
            b"Range: 0 to 100 bar."
        )
        upload = build_upload(
            filename="notes.txt",
            document_format=DocumentFormat.TXT,
            media_type="text/plain",
            content=content,
        )

        parsed = PdfOfficeDocumentParser().parse(upload, content)

        self.assertEqual(
            parsed.parser_name,
            "engineer4me-standard-document-parser",
        )
        self.assertIn("Pressure Transmitter", parsed.full_text)

    def test_legacy_doc_remains_unsupported(self) -> None:
        content = b"legacy word document"
        upload = build_upload(
            filename="manual.doc",
            document_format=DocumentFormat.DOC,
            content=content,
        )

        with self.assertRaises(UnsupportedDocumentFormatError):
            PdfOfficeDocumentParser().parse(upload, content)

    def test_convenience_function_uses_complete_chain(self) -> None:
        content = b"Instrument calibration notes"
        upload = build_upload(
            filename="notes.txt",
            document_format=DocumentFormat.TXT,
            content=content,
        )

        parsed = parse_pdf_office_document(upload, content)

        self.assertEqual(
            parsed.extraction_method,
            ExtractionMethod.NATIVE_TEXT,
        )


class PdfDocumentParserTests(TestCase):
    """Validate native and OCR-assisted PDF parsing."""

    def test_native_pdf_text_metadata_and_safety_block(self) -> None:
        content = build_pdf_bytes(
            "WARNING: Isolate pressure before maintenance.",
            title="Pressure Transmitter Manual",
        )
        upload = build_upload(
            filename="manual.pdf",
            document_format=DocumentFormat.PDF,
            media_type="application/pdf",
            content=content,
        )
        parser = PdfOfficeDocumentParser()

        with patch.object(
            parser.fallback_parser,
            "parse",
        ) as fallback_parse:
            parsed = parser.parse(upload, content)

        fallback_parse.assert_not_called()
        self.assertEqual(
            parsed.title,
            "Pressure Transmitter Manual",
        )
        self.assertEqual(parsed.page_count, 1)
        self.assertEqual(
            parsed.extraction_method,
            ExtractionMethod.NATIVE_TEXT,
        )
        self.assertEqual(
            parsed.pages[0].blocks[0].block_type,
            ContentBlockType.WARNING,
        )
        self.assertEqual(
            parsed.parser_metadata["native_page_count"],
            1,
        )
        self.assertEqual(
            parsed.parser_metadata["ocr_page_count"],
            0,
        )

    def test_native_pdf_pages_have_global_block_sequence(self) -> None:
        content = build_pdf_bytes(
            "First page pressure data.",
            "Second page calibration data.",
        )
        upload = build_upload(
            filename="two-pages.pdf",
            document_format=DocumentFormat.PDF,
            content=content,
        )

        parsed = PdfOfficeDocumentParser().parse(upload, content)

        self.assertEqual(parsed.page_count, 2)
        self.assertEqual(
            [
                block.sequence_number
                for page in parsed.pages
                for block in page.blocks
            ],
            [1, 2],
        )
        self.assertEqual(
            parsed.pages[1].blocks[0].page_number,
            2,
        )

    def test_scanned_pdf_page_uses_ocr_and_remaps_geometry(self) -> None:
        content = build_pdf_bytes("")
        upload = build_upload(
            filename="scan.pdf",
            document_format=DocumentFormat.PDF,
            content=content,
        )
        parser = PdfOfficeDocumentParser()
        ocr_result = build_ocr_result(
            document_id=upload.document_id,
        )

        with patch.object(
            parser.fallback_parser,
            "parse",
            return_value=ocr_result,
        ) as ocr_parse:
            parsed = parser.parse(upload, content)

        image_upload, image_content = ocr_parse.call_args.args
        self.assertEqual(
            image_upload.document_format,
            DocumentFormat.PNG,
        )
        self.assertTrue(image_content.startswith(b"\x89PNG"))
        self.assertEqual(
            parsed.extraction_method,
            ExtractionMethod.OCR,
        )
        self.assertEqual(
            parsed.parser_metadata["ocr_page_count"],
            1,
        )
        block = parsed.pages[0].blocks[0]
        self.assertEqual(block.sequence_number, 1)
        self.assertEqual(block.page_number, 1)
        self.assertAlmostEqual(
            block.bounding_box.width,  # type: ignore[union-attr]
            200.0,
        )
        self.assertAlmostEqual(
            block.bounding_box.height,  # type: ignore[union-attr]
            100.0,
        )

    def test_mixed_native_and_scanned_pdf_is_hybrid(self) -> None:
        content = build_pdf_bytes(
            "Native equipment data.",
            "",
        )
        upload = build_upload(
            filename="hybrid.pdf",
            document_format=DocumentFormat.PDF,
            content=content,
        )
        parser = PdfOfficeDocumentParser()
        ocr_result = build_ocr_result(
            document_id=upload.document_id,
            text="OCR nameplate data.",
        )

        with patch.object(
            parser.fallback_parser,
            "parse",
            return_value=ocr_result,
        ):
            parsed = parser.parse(upload, content)

        self.assertEqual(
            parsed.extraction_method,
            ExtractionMethod.HYBRID,
        )
        self.assertEqual(
            parsed.parser_metadata["native_page_count"],
            1,
        )
        self.assertEqual(
            parsed.parser_metadata["ocr_page_count"],
            1,
        )
        self.assertEqual(
            parsed.pages[1].blocks[0].page_number,
            2,
        )
        self.assertEqual(
            parsed.pages[1].blocks[0].sequence_number,
            2,
        )

    def test_blank_page_warning_is_retained_when_other_page_has_text(
        self,
    ) -> None:
        content = build_pdf_bytes(
            "Readable native page.",
            "",
        )
        upload = build_upload(
            filename="partly-blank.pdf",
            document_format=DocumentFormat.PDF,
            content=content,
        )
        parser = PdfOfficeDocumentParser()

        with patch.object(
            parser.fallback_parser,
            "parse",
            side_effect=EmptyDocumentError("No OCR text."),
        ):
            parsed = parser.parse(upload, content)

        self.assertEqual(
            parsed.parser_metadata["blank_page_count"],
            1,
        )
        self.assertTrue(
            any(
                "No readable native or OCR text" in warning
                for warning in parsed.warnings
            )
        )

    def test_pdf_ocr_failure_identifies_page(self) -> None:
        content = build_pdf_bytes("")
        upload = build_upload(
            filename="failed-scan.pdf",
            document_format=DocumentFormat.PDF,
            content=content,
        )
        parser = PdfOfficeDocumentParser()

        with patch.object(
            parser.fallback_parser,
            "parse",
            side_effect=DocumentParserError("Tesseract failed."),
        ):
            with self.assertRaisesRegex(
                DocumentParserError,
                "OCR failed for PDF page 1",
            ):
                parser.parse(upload, content)

    def test_pdf_with_ocr_disabled_and_no_text_is_empty(self) -> None:
        content = build_pdf_bytes("")
        upload = build_upload(
            filename="blank.pdf",
            document_format=DocumentFormat.PDF,
            content=content,
        )
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                enable_pdf_ocr_fallback=False,
            )
        )

        with self.assertRaises(EmptyDocumentError):
            parser.parse(upload, content)

    def test_pdf_page_limit_is_enforced(self) -> None:
        content = build_pdf_bytes("Page one", "Page two")
        upload = build_upload(
            filename="too-many-pages.pdf",
            document_format=DocumentFormat.PDF,
            content=content,
        )
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                maximum_pdf_pages=1,
            )
        )

        with self.assertRaises(DocumentTooLargeError):
            parser.parse(upload, content)

    def test_pdf_ocr_page_pixel_limit_is_enforced(self) -> None:
        content = build_pdf_bytes("")
        upload = build_upload(
            filename="large-render.pdf",
            document_format=DocumentFormat.PDF,
            content=content,
        )
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                maximum_pdf_ocr_page_pixels=10_000,
            )
        )

        with self.assertRaises(DocumentTooLargeError):
            parser.parse(upload, content)

    def test_pdf_ocr_total_pixel_limit_is_enforced(self) -> None:
        content = build_pdf_bytes("", "")
        upload = build_upload(
            filename="large-total-render.pdf",
            document_format=DocumentFormat.PDF,
            content=content,
        )
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                maximum_pdf_ocr_page_pixels=80_000,
                maximum_pdf_ocr_total_pixels=100_000,
            )
        )
        ocr_result = build_ocr_result(
            document_id=upload.document_id,
        )

        with patch.object(
            parser.fallback_parser,
            "parse",
            return_value=ocr_result,
        ):
            with self.assertRaises(DocumentTooLargeError):
                parser.parse(upload, content)

    def test_pdf_extracted_character_limit_is_enforced(self) -> None:
        content = build_pdf_bytes("More than ten characters")
        upload = build_upload(
            filename="long.pdf",
            document_format=DocumentFormat.PDF,
            content=content,
        )
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                maximum_extracted_characters=10,
            )
        )

        with self.assertRaises(DocumentTooLargeError):
            parser.parse(upload, content)

    def test_pdf_size_mismatch_is_a_warning(self) -> None:
        content = build_pdf_bytes("Readable text")
        upload = build_upload(
            filename="size-mismatch.pdf",
            document_format=DocumentFormat.PDF,
            content=content,
            declared_size=len(content) + 1,
        )

        parsed = PdfOfficeDocumentParser().parse(upload, content)

        self.assertTrue(
            any(
                "size metadata" in warning
                for warning in parsed.warnings
            )
        )

    def test_invalid_pdf_header_and_content_are_rejected(self) -> None:
        parser = PdfOfficeDocumentParser()

        invalid_header = b"not a pdf"
        invalid_header_upload = build_upload(
            filename="invalid.pdf",
            document_format=DocumentFormat.PDF,
            content=invalid_header,
        )

        with self.assertRaises(MalformedDocumentError):
            parser.parse(
                invalid_header_upload,
                invalid_header,
            )

        corrupt_pdf = b"%PDF-1.7\nnot a complete PDF"
        corrupt_upload = build_upload(
            filename="corrupt.pdf",
            document_format=DocumentFormat.PDF,
            content=corrupt_pdf,
        )

        with self.assertRaises(MalformedDocumentError):
            parser.parse(corrupt_upload, corrupt_pdf)


class DocxDocumentParserTests(TestCase):
    """Validate DOCX text, tables, metadata, and limits."""

    def test_docx_structure_safety_table_and_metadata(self) -> None:
        content = build_docx_bytes()
        upload = build_upload(
            filename="valve-manual.docx",
            document_format=DocumentFormat.DOCX,
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            content=content,
        )

        parsed = PdfOfficeDocumentParser().parse(upload, content)

        self.assertEqual(parsed.title, "Valve Maintenance Manual")
        self.assertEqual(parsed.page_count, 1)
        self.assertEqual(
            parsed.extraction_method,
            ExtractionMethod.HYBRID,
        )
        block_types = [
            block.block_type
            for block in parsed.pages[0].blocks
        ]
        self.assertEqual(
            block_types,
            [
                ContentBlockType.TITLE,
                ContentBlockType.HEADING,
                ContentBlockType.DANGER,
                ContentBlockType.LIST,
                ContentBlockType.TABLE,
            ],
        )
        danger_block = parsed.pages[0].blocks[2]
        self.assertNotIn("DANGER:", danger_block.text)
        table_block = parsed.pages[0].blocks[-1]
        self.assertEqual(
            table_block.table.headers,  # type: ignore[union-attr]
            ["Model", "Maximum pressure"],
        )
        self.assertEqual(
            table_block.table.rows[0],  # type: ignore[union-attr]
            ["CV-100", "16 bar"],
        )
        self.assertEqual(
            parsed.parser_metadata["table_count"],
            1,
        )
        self.assertEqual(
            parsed.parser_metadata["document_properties"]["author"],
            "Engineer4Me",
        )

    def test_docx_embedded_image_is_reported(self) -> None:
        content = build_docx_bytes(
            include_table=False,
            include_image=True,
        )
        upload = build_upload(
            filename="illustrated.docx",
            document_format=DocumentFormat.DOCX,
            content=content,
        )

        parsed = PdfOfficeDocumentParser().parse(upload, content)

        self.assertEqual(
            parsed.parser_metadata["embedded_image_count"],
            1,
        )
        self.assertTrue(
            any(
                "embedded DOCX image" in warning
                for warning in parsed.warnings
            )
        )

    def test_empty_docx_is_rejected(self) -> None:
        content = build_docx_bytes(
            include_content=False,
            include_table=False,
        )
        upload = build_upload(
            filename="empty.docx",
            document_format=DocumentFormat.DOCX,
            content=content,
        )

        with self.assertRaises(EmptyDocumentError):
            PdfOfficeDocumentParser().parse(upload, content)

    def test_docx_block_limit_is_enforced(self) -> None:
        content = build_docx_bytes(include_table=False)
        upload = build_upload(
            filename="many-blocks.docx",
            document_format=DocumentFormat.DOCX,
            content=content,
        )
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                maximum_docx_blocks=1,
            )
        )

        with self.assertRaises(DocumentTooLargeError):
            parser.parse(upload, content)

    def test_docx_table_row_limit_is_enforced(self) -> None:
        content = build_docx_bytes()
        upload = build_upload(
            filename="large-table.docx",
            document_format=DocumentFormat.DOCX,
            content=content,
        )
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                maximum_table_rows=2,
            )
        )

        with self.assertRaises(DocumentTooLargeError):
            parser.parse(upload, content)

    def test_docx_cell_character_limit_is_enforced(self) -> None:
        document = Document()
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "12345"
        output = BytesIO()
        document.save(output)
        content = output.getvalue()
        upload = build_upload(
            filename="long-cell.docx",
            document_format=DocumentFormat.DOCX,
            content=content,
        )
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                maximum_cell_characters=4,
            )
        )

        with self.assertRaises(DocumentTooLargeError):
            parser.parse(upload, content)

    def test_malformed_docx_is_rejected(self) -> None:
        content = b"PK\x03\x04not a valid OOXML archive"
        upload = build_upload(
            filename="invalid.docx",
            document_format=DocumentFormat.DOCX,
            content=content,
        )

        with self.assertRaises(MalformedDocumentError):
            PdfOfficeDocumentParser().parse(upload, content)


class SpreadsheetDocumentParserTests(TestCase):
    """Validate XLSX and legacy XLS spreadsheet extraction."""

    def test_xlsx_values_range_and_metadata_are_extracted(self) -> None:
        content = build_xlsx_bytes()
        upload = build_upload(
            filename="catalogue.xlsx",
            document_format=DocumentFormat.XLSX,
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "spreadsheetml.sheet"
            ),
            content=content,
        )

        parsed = PdfOfficeDocumentParser().parse(upload, content)

        self.assertEqual(parsed.title, "Instrument Catalogue")
        self.assertEqual(
            parsed.extraction_method,
            ExtractionMethod.TABLE_EXTRACTION,
        )
        table_block = parsed.pages[0].blocks[0]
        spreadsheet_range = table_block.spreadsheet_range
        self.assertIsNotNone(spreadsheet_range)
        assert spreadsheet_range is not None
        self.assertEqual(
            spreadsheet_range.start_cell,
            "A1",
        )
        self.assertEqual(
            spreadsheet_range.end_cell,
            "D2",
        )
        self.assertEqual(
            table_block.table.headers,  # type: ignore[union-attr]
            [
                "Model",
                "Range",
                "Enabled",
                "Calibration date",
            ],
        )
        self.assertEqual(
            table_block.table.rows[0],  # type: ignore[union-attr]
            [
                "PT-100",
                "100",
                "TRUE",
                "2026-01-02 03:04:05",
            ],
        )
        self.assertEqual(
            parsed.parser_metadata["sheet_count"],
            1,
        )
        self.assertEqual(
            parsed.parser_metadata["formula_values"],
            "cached",
        )

    def test_xlsx_blank_sheet_is_retained_with_warning(self) -> None:
        content = build_xlsx_bytes(include_blank_sheet=True)
        upload = build_upload(
            filename="two-sheets.xlsx",
            document_format=DocumentFormat.XLSX,
            content=content,
        )

        parsed = PdfOfficeDocumentParser().parse(upload, content)

        self.assertEqual(parsed.page_count, 2)
        self.assertEqual(
            parsed.parser_metadata["non_empty_sheet_count"],
            1,
        )
        self.assertTrue(
            any(
                "Archived" in warning
                for warning in parsed.warnings
            )
        )

    def test_empty_xlsx_is_rejected(self) -> None:
        content = build_xlsx_bytes(include_values=False)
        upload = build_upload(
            filename="empty.xlsx",
            document_format=DocumentFormat.XLSX,
            content=content,
        )

        with self.assertRaises(EmptyDocumentError):
            PdfOfficeDocumentParser().parse(upload, content)

    def test_xlsx_sheet_limit_is_enforced(self) -> None:
        content = build_xlsx_bytes(include_blank_sheet=True)
        upload = build_upload(
            filename="many-sheets.xlsx",
            document_format=DocumentFormat.XLSX,
            content=content,
        )
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                maximum_workbook_sheets=1,
            )
        )

        with self.assertRaises(DocumentTooLargeError):
            parser.parse(upload, content)

    def test_xlsx_row_column_and_cell_limits_are_enforced(
        self,
    ) -> None:
        content = build_xlsx_bytes()
        upload = build_upload(
            filename="bounded.xlsx",
            document_format=DocumentFormat.XLSX,
            content=content,
        )

        for config in (
            PdfOfficeDocumentParserConfig(
                maximum_spreadsheet_rows=1,
            ),
            PdfOfficeDocumentParserConfig(
                maximum_spreadsheet_columns=3,
            ),
            PdfOfficeDocumentParserConfig(
                maximum_spreadsheet_cells=7,
            ),
        ):
            with self.subTest(config=config):
                with self.assertRaises(DocumentTooLargeError):
                    PdfOfficeDocumentParser(config).parse(
                        upload,
                        content,
                    )

    def test_xlsx_cell_character_limit_is_enforced(self) -> None:
        content = build_xlsx_bytes(long_value="TOO-LONG")
        upload = build_upload(
            filename="long-cell.xlsx",
            document_format=DocumentFormat.XLSX,
            content=content,
        )
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                maximum_cell_characters=7,
            )
        )

        with self.assertRaises(DocumentTooLargeError):
            parser.parse(upload, content)

    def test_legacy_xls_common_cell_types_are_extracted(self) -> None:
        content = _OLE_SIGNATURE + b"mock-xls-content"
        upload = build_upload(
            filename="legacy.xls",
            document_format=DocumentFormat.XLS,
            media_type="application/vnd.ms-excel",
            content=content,
        )
        workbook = build_fake_xls_workbook()

        with patch(
            "app.ingestion.pdf_office_document_parser."
            "xlrd.open_workbook",
            return_value=workbook,
        ) as open_workbook:
            parsed = PdfOfficeDocumentParser().parse(
                upload,
                content,
        )

        open_workbook.assert_called_once_with(
            file_contents=content,
            on_demand=True,
        )
        self.assertTrue(workbook.resources_released)
        table = parsed.pages[0].blocks[0].table
        self.assertIsNotNone(table)
        assert table is not None
        row = table.rows[0]
        self.assertEqual(row[0], "LT-100")
        self.assertEqual(row[1], "25")
        self.assertEqual(row[2], "TRUE")
        self.assertTrue(row[3].startswith("1900-01-01"))
        self.assertEqual(row[4], "#DIV/0!")
        self.assertEqual(
            parsed.parser_metadata["source_format"],
            DocumentFormat.XLS.value,
        )

    def test_legacy_xls_sheet_limit_releases_resources(self) -> None:
        content = _OLE_SIGNATURE + b"mock-xls-content"
        upload = build_upload(
            filename="many-sheets.xls",
            document_format=DocumentFormat.XLS,
            content=content,
        )
        workbook = FakeXlsWorkbook(
            [
                FakeXlsSheet("One", []),
                FakeXlsSheet("Two", []),
            ]
        )
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                maximum_workbook_sheets=1,
            )
        )

        with patch(
            "app.ingestion.pdf_office_document_parser."
            "xlrd.open_workbook",
            return_value=workbook,
        ):
            with self.assertRaises(DocumentTooLargeError):
                parser.parse(upload, content)

        self.assertTrue(workbook.resources_released)

    def test_legacy_xls_invalid_header_is_rejected(self) -> None:
        content = b"not an OLE workbook"
        upload = build_upload(
            filename="invalid.xls",
            document_format=DocumentFormat.XLS,
            content=content,
        )

        with self.assertRaises(MalformedDocumentError):
            PdfOfficeDocumentParser().parse(upload, content)

    def test_legacy_xls_library_error_is_normalised(self) -> None:
        content = _OLE_SIGNATURE + b"corrupt"
        upload = build_upload(
            filename="corrupt.xls",
            document_format=DocumentFormat.XLS,
            content=content,
        )

        with patch(
            "app.ingestion.pdf_office_document_parser."
            "xlrd.open_workbook",
            side_effect=xlrd.XLRDError("invalid workbook"),
        ):
            with self.assertRaises(MalformedDocumentError):
                PdfOfficeDocumentParser().parse(upload, content)


class OfficeArchiveSecurityTests(TestCase):
    """Validate OOXML archive and common input safeguards."""

    def test_password_protected_upload_is_rejected(self) -> None:
        content = build_pdf_bytes("Protected content")
        upload = build_upload(
            filename="protected.pdf",
            document_format=DocumentFormat.PDF,
            content=content,
            password_protected=True,
        )

        with self.assertRaises(PasswordProtectedDocumentError):
            PdfOfficeDocumentParser().parse(upload, content)

    def test_empty_non_bytes_and_oversized_input_are_rejected(
        self,
    ) -> None:
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                maximum_document_bytes=4,
            )
        )

        empty_upload = build_upload(
            filename="empty.pdf",
            document_format=DocumentFormat.PDF,
            content=b"",
        )

        with self.assertRaises(EmptyDocumentError):
            parser.parse(empty_upload, b"")

        type_upload = build_upload(
            filename="type.pdf",
            document_format=DocumentFormat.PDF,
            content=b"data",
        )

        with self.assertRaises(TypeError):
            parser.parse(
                type_upload,
                "data",  # type: ignore[arg-type]
            )

        large_content = b"12345"
        large_upload = build_upload(
            filename="large.pdf",
            document_format=DocumentFormat.PDF,
            content=large_content,
        )

        with self.assertRaises(DocumentTooLargeError):
            parser.parse(large_upload, large_content)

    def test_archive_member_limit_is_enforced(self) -> None:
        content = build_xlsx_bytes()
        upload = build_upload(
            filename="members.xlsx",
            document_format=DocumentFormat.XLSX,
            content=content,
        )
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                maximum_archive_members=1,
            )
        )

        with self.assertRaises(DocumentTooLargeError):
            parser.parse(upload, content)

    def test_archive_uncompressed_size_limit_is_enforced(self) -> None:
        content = build_docx_bytes()
        upload = build_upload(
            filename="expanded.docx",
            document_format=DocumentFormat.DOCX,
            content=content,
        )
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                maximum_archive_uncompressed_bytes=100,
            )
        )

        with self.assertRaises(DocumentTooLargeError):
            parser.parse(upload, content)

    def test_archive_compression_ratio_limit_is_enforced(self) -> None:
        content = append_zip_member(
            build_xlsx_bytes(),
            member_name="xl/repeated.bin",
            member_content=b"A" * 20_000,
        )
        upload = build_upload(
            filename="compressed.xlsx",
            document_format=DocumentFormat.XLSX,
            content=content,
        )
        parser = PdfOfficeDocumentParser(
            PdfOfficeDocumentParserConfig(
                maximum_archive_compression_ratio=10.0,
            )
        )

        with self.assertRaises(UnsafeOfficeArchiveError):
            parser.parse(upload, content)

    def test_archive_path_traversal_is_rejected(self) -> None:
        content = append_zip_member(
            build_docx_bytes(),
            member_name="../escape.xml",
            member_content=b"<root />",
        )
        upload = build_upload(
            filename="traversal.docx",
            document_format=DocumentFormat.DOCX,
            content=content,
        )

        with self.assertRaises(UnsafeOfficeArchiveError):
            PdfOfficeDocumentParser().parse(upload, content)

    def test_vba_macro_member_is_rejected(self) -> None:
        content = append_zip_member(
            build_xlsx_bytes(),
            member_name="xl/vbaProject.bin",
            member_content=b"macro",
        )
        upload = build_upload(
            filename="macro.xlsx",
            document_format=DocumentFormat.XLSX,
            content=content,
        )

        with self.assertRaises(UnsafeOfficeArchiveError):
            PdfOfficeDocumentParser().parse(upload, content)

    def test_unsafe_xml_declarations_are_rejected(self) -> None:
        unsafe_xml = (
            b'<?xml version="1.0"?>'
            b'<!DOCTYPE root [<!ENTITY secret "unsafe">]>'
            b"<root>&secret;</root>"
        )
        content = append_zip_member(
            build_docx_bytes(),
            member_name="customXml/unsafe.xml",
            member_content=unsafe_xml,
        )
        upload = build_upload(
            filename="unsafe-xml.docx",
            document_format=DocumentFormat.DOCX,
            content=content,
        )

        with self.assertRaises(UnsafeOfficeArchiveError):
            PdfOfficeDocumentParser().parse(upload, content)

    def test_duplicate_archive_member_is_rejected(self) -> None:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            content = append_zip_member(
                build_xlsx_bytes(),
                member_name="[Content_Types].xml",
                member_content=b"<Types />",
            )

        upload = build_upload(
            filename="duplicate.xlsx",
            document_format=DocumentFormat.XLSX,
            content=content,
        )

        with self.assertRaises(UnsafeOfficeArchiveError):
            PdfOfficeDocumentParser().parse(upload, content)

    def test_missing_required_ooxml_member_is_rejected(self) -> None:
        output = BytesIO()

        with ZipFile(
            output,
            mode="w",
            compression=ZIP_DEFLATED,
        ) as archive:
            archive.writestr(
                "[Content_Types].xml",
                (
                    b'<?xml version="1.0"?>'
                    b'<Types xmlns="http://schemas.openxmlformats.org/'
                    b'package/2006/content-types" />'
                ),
            )

        content = output.getvalue()
        upload = build_upload(
            filename="missing-document.docx",
            document_format=DocumentFormat.DOCX,
            content=content,
        )

        with self.assertRaises(MalformedDocumentError):
            PdfOfficeDocumentParser().parse(upload, content)


if __name__ == "__main__":
    import unittest

    unittest.main()
