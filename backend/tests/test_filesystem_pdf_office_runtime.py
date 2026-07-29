"""Integration tests for PDF and Office parser runtime composition."""

from __future__ import annotations

from dataclasses import FrozenInstanceError
from hashlib import sha256
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest import TestCase

from docx import Document

from app.ingestion.document_models import (
    ContentBlockType,
    DocumentFormat,
    DocumentSource,
    DocumentUpload,
    ExtractionMethod,
)
from app.ingestion.document_parser import (
    DocumentTooLargeError,
)
from app.ingestion.document_processing_orchestrator import (
    DocumentProcessingOrchestratorConfig,
)
from app.ingestion.filesystem_document_content_loader import (
    FilesystemDocumentContentLoaderConfig,
)
from app.ingestion.filesystem_document_processing_runtime import (
    FilesystemDocumentProcessingRuntime,
    FilesystemDocumentProcessingRuntimeConfig,
)
from app.ingestion.ingestion_job_service import IngestionJobService
from app.ingestion.pdf_office_document_parser import (
    PasswordProtectedDocumentError,
    PdfOfficeDocumentParser,
    PdfOfficeDocumentParserConfig,
)


def build_upload(
    *,
    filename: str,
    document_format: DocumentFormat,
    content: bytes,
    media_type: str | None = None,
    password_protected: bool = False,
) -> DocumentUpload:
    """Build valid upload metadata for one runtime integration test."""

    return DocumentUpload(
        filename=filename,
        document_format=document_format,
        media_type=media_type,
        size_bytes=len(content),
        storage_key=f"runtime-tests/{filename}",
        checksum_sha256=sha256(content).hexdigest(),
        source=DocumentSource(
            source_name="Filesystem runtime integration test",
        ),
        password_protected=password_protected,
    )


def build_docx_bytes() -> bytes:
    """Build one small DOCX document with text and a table."""

    document = Document()
    document.core_properties.title = "Runtime Valve Manual"
    document.add_heading("Runtime Valve Manual", level=0)
    document.add_paragraph(
        "WARNING: Isolate and depressurise before maintenance."
    )
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Model"
    table.cell(0, 1).text = "Pressure"
    table.cell(1, 0).text = "CV-100"
    table.cell(1, 1).text = "16 bar"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


class FilesystemPdfOfficeRuntimeTests(TestCase):
    """Validate production composition of the complete parser chain."""

    def setUp(self) -> None:
        """Create one isolated filesystem root for each runtime."""

        self._temporary_directory = TemporaryDirectory()
        self.root_directory = Path(
            self._temporary_directory.name
        )

    def tearDown(self) -> None:
        """Remove the isolated filesystem root."""

        self._temporary_directory.cleanup()

    def build_runtime(
        self,
        *,
        parser_config: (
            PdfOfficeDocumentParserConfig | None
        ) = None,
    ) -> FilesystemDocumentProcessingRuntime:
        """Build the concrete runtime with optional parser controls."""

        config_values = {
            "content_loader_config": (
                FilesystemDocumentContentLoaderConfig(
                    root_directory=self.root_directory,
                )
            ),
            "orchestrator": (
                DocumentProcessingOrchestratorConfig()
            ),
        }

        if parser_config is not None:
            config_values["document_parser"] = parser_config

        config = FilesystemDocumentProcessingRuntimeConfig(
            **config_values,
        )

        return FilesystemDocumentProcessingRuntime(
            job_service=IngestionJobService(),
            config=config,
        )

    def test_runtime_config_defaults_are_bounded_and_frozen(
        self,
    ) -> None:
        """Runtime config owns immutable complete-parser controls."""

        config = FilesystemDocumentProcessingRuntimeConfig(
            content_loader_config=(
                FilesystemDocumentContentLoaderConfig(
                    root_directory=self.root_directory,
                )
            ),
        )

        self.assertIsInstance(
            config.document_parser,
            PdfOfficeDocumentParserConfig,
        )
        self.assertEqual(
            config.document_parser.maximum_document_bytes,
            25 * 1024 * 1024,
        )
        self.assertEqual(
            config.document_parser.maximum_pdf_pages,
            500,
        )
        self.assertTrue(
            config.document_parser.enable_pdf_ocr_fallback
        )

        with self.assertRaises(FrozenInstanceError):
            config.document_parser = (  # type: ignore[misc]
                PdfOfficeDocumentParserConfig()
            )

    def test_runtime_config_rejects_invalid_parser_controls(
        self,
    ) -> None:
        """Only validated PDF and Office parser config is accepted."""

        with self.assertRaisesRegex(
            TypeError,
            "PdfOfficeDocumentParserConfig",
        ):
            FilesystemDocumentProcessingRuntimeConfig(
                content_loader_config=(
                    FilesystemDocumentContentLoaderConfig(
                        root_directory=self.root_directory,
                    )
                ),
                document_parser=object(),  # type: ignore[arg-type]
            )

    def test_runtime_injects_one_shared_complete_parser(
        self,
    ) -> None:
        """Runtime and orchestrator share one configured parser."""

        parser_config = PdfOfficeDocumentParserConfig(
            maximum_pdf_pages=25,
            maximum_workbook_sheets=10,
        )
        runtime = self.build_runtime(
            parser_config=parser_config,
        )

        self.assertIsInstance(
            runtime.document_parser,
            PdfOfficeDocumentParser,
        )
        self.assertIs(
            runtime.document_parser.config,
            parser_config,
        )
        self.assertIs(
            runtime.orchestrator._parser,
            runtime.document_parser,
        )
        self.assertEqual(
            runtime.document_parser.config.maximum_pdf_pages,
            25,
        )
        self.assertEqual(
            runtime.document_parser.config.maximum_workbook_sheets,
            10,
        )

    def test_runtime_parser_reports_complete_format_chain(
        self,
    ) -> None:
        """Production parser supports PDF, Office, text, and images."""

        runtime = self.build_runtime()

        for filename, document_format in (
            ("manual.pdf", DocumentFormat.PDF),
            ("manual.docx", DocumentFormat.DOCX),
            ("catalogue.xlsx", DocumentFormat.XLSX),
            ("legacy.xls", DocumentFormat.XLS),
            ("notes.txt", DocumentFormat.TXT),
            ("table.csv", DocumentFormat.CSV),
            ("nameplate.png", DocumentFormat.PNG),
        ):
            with self.subTest(document_format=document_format):
                content = b"placeholder"
                upload = build_upload(
                    filename=filename,
                    document_format=document_format,
                    content=content,
                )
                self.assertTrue(
                    runtime.document_parser.supports(upload)
                )

    def test_runtime_parser_delegates_existing_text_formats(
        self,
    ) -> None:
        """Text documents still use the established native parser."""

        content = (
            b"# Pressure Transmitter\n\n"
            b"Operating range: 0 to 100 bar."
        )
        upload = build_upload(
            filename="pressure-notes.txt",
            document_format=DocumentFormat.TXT,
            media_type="text/plain",
            content=content,
        )

        parsed = self.build_runtime().document_parser.parse(
            upload,
            content,
        )

        self.assertEqual(
            parsed.parser_name,
            "engineer4me-standard-document-parser",
        )
        self.assertEqual(
            parsed.extraction_method,
            ExtractionMethod.NATIVE_TEXT,
        )
        self.assertIn(
            "Operating range: 0 to 100 bar.",
            parsed.full_text,
        )

    def test_runtime_parser_processes_docx_text_and_tables(
        self,
    ) -> None:
        """The production runtime parses DOCX structured content."""

        content = build_docx_bytes()
        upload = build_upload(
            filename="valve-manual.docx",
            document_format=DocumentFormat.DOCX,
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            content=content,
        )

        parsed = self.build_runtime().document_parser.parse(
            upload,
            content,
        )

        self.assertEqual(parsed.title, "Runtime Valve Manual")
        self.assertEqual(
            parsed.extraction_method,
            ExtractionMethod.HYBRID,
        )
        self.assertEqual(parsed.page_count, 1)
        self.assertEqual(
            [
                block.block_type
                for block in parsed.pages[0].blocks
            ],
            [
                ContentBlockType.TITLE,
                ContentBlockType.WARNING,
                ContentBlockType.TABLE,
            ],
        )
        table = parsed.pages[0].blocks[-1].table
        self.assertIsNotNone(table)
        assert table is not None
        self.assertEqual(
            table.headers,
            ["Model", "Pressure"],
        )
        self.assertEqual(
            table.rows,
            [["CV-100", "16 bar"]],
        )

    def test_runtime_parser_enforces_custom_size_limit(
        self,
    ) -> None:
        """Configured resource limits survive runtime composition."""

        parser_config = PdfOfficeDocumentParserConfig(
            maximum_document_bytes=4,
        )
        runtime = self.build_runtime(
            parser_config=parser_config,
        )
        content = b"12345"
        upload = build_upload(
            filename="oversized.pdf",
            document_format=DocumentFormat.PDF,
            content=content,
        )

        with self.assertRaises(DocumentTooLargeError):
            runtime.document_parser.parse(upload, content)

    def test_runtime_parser_rejects_password_protected_input(
        self,
    ) -> None:
        """Password-protection policy is active in production wiring."""

        content = b"%PDF-1.7\nprotected"
        upload = build_upload(
            filename="protected.pdf",
            document_format=DocumentFormat.PDF,
            content=content,
            password_protected=True,
        )

        with self.assertRaises(PasswordProtectedDocumentError):
            self.build_runtime().document_parser.parse(
                upload,
                content,
            )


if __name__ == "__main__":
    import unittest

    unittest.main()
